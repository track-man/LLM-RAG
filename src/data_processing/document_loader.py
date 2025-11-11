"""
文档加载模块
负责从指定目录加载原始文档，支持多种文件格式
"""
import os
import logging
import csv
from typing import List, Dict, Any
from pathlib import Path
import config

# 配置日志
logger = logging.getLogger(__name__)

class DocumentLoader:
    """文档加载器类"""
    
    def __init__(self):
        """初始化文档加载器"""
        # 支持的文件扩展名
        self.supported_extensions = {'.txt', '.md', '.pdf', '.docx', '.doc', '.csv'}
        
        # 文件编码格式
        self.encodings = ['utf-8', 'gbk', 'gb2312', 'latin1']
    
    def load_documents(self, doc_dir: str) -> List[Dict[str, Any]]:
        """
        从指定目录加载所有支持的文档
        
        Args:
            doc_dir: 文档目录路径
            
        Returns:
            文档列表，每个元素为 {"text": 文本内容, "metadata": 元数据}
        """
        if not os.path.exists(doc_dir):
            raise ValueError(f"文档目录不存在: {doc_dir}")
        
        logger.info(f"开始加载文档，目录: {doc_dir}")
        
        documents = []
        doc_path = Path(doc_dir)
        
        # 递归遍历所有文件
        for file_path in doc_path.rglob('*'):
            if file_path.is_file():
                # 检查文件扩展名
                if file_path.suffix.lower() not in self.supported_extensions:
                    continue
                
                try:
                    # 加载文档内容
                    doc_content = self._load_single_document(file_path)
                    if doc_content:
                        # 构建文档对象
                        document = {
                            "text": doc_content,
                            "metadata": {
                                "source": str(file_path),
                            }
                        }
                        documents.append(document)
                        logger.debug(f"成功加载文档: {file_path.name}")
                    
                except Exception as e:
                    logger.warning(f"加载文档失败 {file_path.name}: {str(e)}")
                    continue
        
        logger.info(f"文档加载完成，共加载 {len(documents)} 个文档")
        return documents
    
    def _load_single_document(self, file_path: Path) -> str:
        """
        加载单个文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档文本内容
        """
        content = ""
        
        # 根据文件类型选择加载方法
        if file_path.suffix.lower() == '.pdf':
            content = self._load_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            content = self._load_word_document(file_path)
        elif file_path.suffix.lower() == '.csv':
            content = self._load_csv(file_path)
        else:
            # 文本文件
            content = self._load_text_file(file_path)
        
        return content
    
    def _load_csv(self, file_path: Path) -> str:
        """
        加载csv文件
        
        Args:
            file_path: csv文件路径
            
        Returns:
            csv文件文本内容
        """
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as csvfile:
                    # 读取CSV内容
                    csv_reader = csv.reader(csvfile)
                    rows = list(csv_reader)
                    
                    # 将内容转换为字符串
                    result_lines = []
                    for row in rows:
                        # 将每行转换为逗号分隔的字符串
                        result_lines.append(','.join(str(cell) for cell in row))
                    
                    return '\n'.join(result_lines)
                    
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"加载csv文档失败 {file_path.name}: {str(e)}")
        return ""
    
    def _load_text_file(self, file_path: Path) -> str:
        """
        加载文本文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件文本内容
        """
        # 尝试不同的编码格式
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    logger.debug(f"使用编码 {encoding} 成功读取文件: {file_path.name}")
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"读取文件失败 {file_path.name}: {str(e)}")
                break
        
        # 如果所有编码都失败，尝试二进制模式读取并忽略错误
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                logger.warning(f"使用忽略错误模式读取文件: {file_path.name}")
                return content
        except Exception as e:
            logger.error(f"无法读取文件 {file_path.name}: {str(e)}")
            return ""
    
    def _load_pdf(self, file_path: Path) -> str:
        """
        加载PDF文件（需要安装PyPDF2或pdfplumber）
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            PDF文本内容
        """
        try:
            # 尝试使用pdfplumber
            import pdfplumber
            
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            content = '\n'.join(text_content)
            logger.debug(f"成功加载PDF文件: {file_path.name}")
            return content
            
        except ImportError:
            logger.warning("pdfplumber未安装，无法处理PDF文件")
            return ""
        except Exception as e:
            logger.warning(f"加载PDF文件失败 {file_path.name}: {str(e)}")
            return ""
    
    def _load_word_document(self, file_path: Path) -> str:
        """
        加载Word文档（需要安装python-docx）
        
        Args:
            file_path: Word文件路径
            
        Returns:
            Word文档文本内容
        """
        try:
            # 尝试使用python-docx
            from docx import Document
            
            doc = Document(str(file_path))
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            content = '\n'.join(text_content)
            logger.debug(f"成功加载Word文档: {file_path.name}")
            return content
            
        except ImportError:
            logger.warning("python-docx未安装，无法处理Word文档")
            return ""
        except Exception as e:
            logger.warning(f"加载Word文档失败 {file_path.name}: {str(e)}")
            return ""
    
    def filter_documents(
        self, 
        documents: List[Dict[str, Any]], 
        min_length: int = 10,
        max_length: int = 100000
    ) -> List[Dict[str, Any]]:
        """
        过滤文档（根据长度等条件）
        
        Args:
            documents: 文档列表
            min_length: 最小文本长度
            max_length: 最大文本长度
            
        Returns:
            过滤后的文档列表
        """
        filtered_docs = []
        
        for doc in documents:
            text_length = len(doc["text"].strip())
            
            if min_length <= text_length <= max_length:
                filtered_docs.append(doc)
            else:
                logger.debug(f"过滤掉文档 {doc['metadata']['filename']}: 长度 {text_length} 不在范围内 [{min_length}, {max_length}]")
        
        logger.info(f"文档过滤完成: {len(documents)} -> {len(filtered_docs)} 个文档")
        return filtered_docs
    
    def get_document_stats(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        获取文档统计信息
        
        Args:
            documents: 文档列表
            
        Returns:
            统计信息字典
        """
        if not documents:
            return {"total_documents": 0}
        
        total_chars = sum(len(doc["text"]) for doc in documents)
        total_words = sum(len(doc["text"].split()) for doc in documents)
        
        # 文件类型统计
        file_types = {}
        for doc in documents:
            ext = doc["metadata"]["file_extension"]
            file_types[ext] = file_types.get(ext, 0) + 1
        
        stats = {
            "total_documents": len(documents),
            "total_characters": total_chars,
            "total_words": total_words,
            "avg_characters_per_doc": total_chars / len(documents),
            "avg_words_per_doc": total_words / len(documents),
            "file_types": file_types,
            "largest_document": max(documents, key=lambda x: len(x["text"]))["metadata"]["filename"],
            "smallest_document": min(documents, key=lambda x: len(x["text"]))["metadata"]["filename"]
        }
        
        logger.info(f"文档统计: {stats}")
        return stats


# 全局文档加载器实例
_document_loader = None

def get_document_loader() -> DocumentLoader:
    """获取全局文档加载器实例（单例模式）"""
    global _document_loader
    if _document_loader is None:
        _document_loader = DocumentLoader()
    return _document_loader

def load_documents(doc_dir: str) -> List[Dict[str, Any]]:
    """
    加载文档的便捷函数
    
    Args:
        doc_dir: 文档目录路径
        
    Returns:
        文档列表，每个元素为 {"text": 文本内容, "metadata": 元数据}
    """
    loader = get_document_loader()
    return loader.load_documents(doc_dir)

def filter_documents(
    documents: List[Dict[str, Any]], 
    min_length: int = 10,
    max_length: int = 100000
) -> List[Dict[str, Any]]:
    """
    过滤文档的便捷函数
    
    Args:
        documents: 文档列表
        min_length: 最小文本长度
        max_length: 最大文本长度
        
    Returns:
        过滤后的文档列表
    """
    loader = get_document_loader()
    return loader.filter_documents(documents, min_length, max_length)

def get_document_stats(documents: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    获取文档统计信息的便捷函数
    
    Args:
        documents: 文档列表
        
    Returns:
        统计信息字典
    """
    loader = get_document_loader()
    return loader.get_document_stats(documents)


if __name__ == "__main__":
    # 测试代码
    import sys
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
    try:
        # 测试文档加载
        print("正在测试文档加载...")
        documents = load_documents(config.RAW_DOC_DIR)
        print(f"成功加载 {len(documents)} 个文档")
        
        if documents:
            # 显示统计信息
            stats = get_document_stats(documents)
            print(f"文档统计: {stats}")
            
            # 显示第一个文档的元数据
            print(f"第一个文档元数据: {documents[0]['metadata']}")
            print(f"第一个文档前100个字符: {documents[0]['text'][:100]}...")
        
    except Exception as e:
        print(f"测试失败: {str(e)}")